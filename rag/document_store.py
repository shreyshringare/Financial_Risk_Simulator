"""
Session-scoped document store for user-uploaded files.
Uses the same ChromaDB instance and embeddings as the financial knowledge base,
but in a separate collection ('session_docs') with session_id metadata filtering.

Supported formats: PDF, DOCX, TXT, MD, CSV.
"""
import io
import csv
from pathlib import Path
from typing import List

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document

PERSIST_DIRECTORY = "./chroma_db"
SESSION_COLLECTION = "session_docs"
CHUNK_SIZE = 600
CHUNK_OVERLAP = 60

_EMBEDDINGS = None


def _get_embeddings() -> HuggingFaceEmbeddings:
    global _EMBEDDINGS
    if _EMBEDDINGS is None:
        _EMBEDDINGS = HuggingFaceEmbeddings(
            model_name="BAAI/bge-base-en-v1.5",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
    return _EMBEDDINGS


def _get_vectorstore() -> Chroma:
    return Chroma(
        collection_name=SESSION_COLLECTION,
        embedding_function=_get_embeddings(),
        persist_directory=PERSIST_DIRECTORY,
    )


def parse_document(filename: str, content: bytes) -> str:
    """
    Extract plain text from uploaded file.
    Supports: .pdf, .docx, .txt, .md, .csv
    Raises ValueError on unsupported type or parse failure.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(p for p in pages if p.strip())
        except Exception as e:
            raise ValueError(f"PDF parse error: {e}") from e

    elif suffix == ".docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise ValueError(f"DOCX parse error: {e}") from e

    elif suffix in (".txt", ".md"):
        try:
            return content.decode("utf-8", errors="replace")
        except Exception as e:
            raise ValueError(f"Text decode error: {e}") from e

    elif suffix == ".csv":
        try:
            text = content.decode("utf-8", errors="replace")
            reader = csv.reader(io.StringIO(text))
            rows = [", ".join(row) for row in reader]
            return "\n".join(rows)
        except Exception as e:
            raise ValueError(f"CSV parse error: {e}") from e

    else:
        raise ValueError(f"Unsupported file type: '{suffix}'. Allowed: pdf, docx, txt, md, csv.")


def ingest_document(session_id: str, filename: str, content: bytes) -> int:
    """
    Parse, chunk, embed, and store a document for a session.
    Returns number of chunks stored.
    """
    text = parse_document(filename, content)
    if not text.strip():
        raise ValueError("Document is empty or contains no extractable text.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_text(text)

    docs = [
        Document(
            page_content=chunk,
            metadata={"session_id": session_id, "filename": filename, "chunk_index": i},
        )
        for i, chunk in enumerate(chunks)
    ]

    vs = _get_vectorstore()
    vs.add_documents(docs)
    return len(docs)


def query_session_docs(session_id: str, query: str, k: int = 3) -> str:
    """
    Retrieve top-k chunks for this session most relevant to the query.
    Returns formatted string, empty string if no docs found.
    """
    vs = _get_vectorstore()
    results = vs.similarity_search(
        query,
        k=k,
        filter={"session_id": session_id},
    )
    if not results:
        return ""
    return "\n\n---\n\n".join(doc.page_content for doc in results)


def delete_session_docs(session_id: str) -> None:
    """Remove all chunks for a session from the store."""
    vs = _get_vectorstore()
    collection = vs._collection  # direct chromadb collection access
    ids = collection.get(where={"session_id": session_id})["ids"]
    if ids:
        collection.delete(ids=ids)


def list_session_files(session_id: str) -> List[str]:
    """Return unique filenames uploaded for this session."""
    vs = _get_vectorstore()
    collection = vs._collection
    result = collection.get(where={"session_id": session_id}, include=["metadatas"])
    filenames = list({m["filename"] for m in result["metadatas"]})
    return sorted(filenames)
